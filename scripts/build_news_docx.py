#!/usr/bin/env python
"""Build a Chinese research news release DOCX from a JSON spec."""

from __future__ import annotations

import argparse
import json
from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


DEFAULT_IMAGE_WIDTH_CM = 14.65
PUBLICATION_FIELDS = [
    ("project_support", "项目支持"),
    ("paper_link", "论文链接"),
    ("online_publication_date", "在线发表日期"),
    ("impact_factor", "影响因子"),
]


def configure_document(doc: Document) -> None:
    for section in doc.sections:
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(3.17)
        section.right_margin = Cm(3.17)
    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    normal.font.size = Pt(12)


def clear_document(doc: Document) -> None:
    body = doc._body  # python-docx does not expose a public clear API.
    if hasattr(body, "clear_content"):
        body.clear_content()


def add_text_paragraph(doc: Document, text: str, *, align: int | None = None, size_pt: float | None = None, bold: bool = False) -> None:
    paragraph = doc.add_paragraph()
    if align is not None:
        paragraph.alignment = align
    run = paragraph.add_run(text)
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    if size_pt:
        run.font.size = Pt(size_pt)
    if bold:
        run.bold = True


def add_figure(doc: Document, figure: dict[str, Any], base_dir: Path) -> None:
    image = figure.get("image") or figure.get("path") or ""
    caption = figure.get("caption") or figure.get("title") or ""
    note = figure.get("note") or figure.get("description") or ""
    width_cm = float(figure.get("width_cm") or DEFAULT_IMAGE_WIDTH_CM)

    if image:
        image_path = Path(image)
        if not image_path.is_absolute():
            image_path = base_dir / image_path
        if image_path.exists():
            paragraph = doc.add_paragraph()
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = paragraph.add_run()
            run.add_picture(str(image_path), width=Cm(width_cm))
        else:
            add_text_paragraph(doc, f"【待插图：{image_path}】", align=WD_ALIGN_PARAGRAPH.CENTER)
    else:
        add_text_paragraph(doc, "【待插图】", align=WD_ALIGN_PARAGRAPH.CENTER)

    if caption:
        add_text_paragraph(doc, caption, align=WD_ALIGN_PARAGRAPH.CENTER)
    if note:
        add_text_paragraph(doc, note, align=WD_ALIGN_PARAGRAPH.CENTER)


def add_publication_info(doc: Document, spec: dict[str, Any]) -> None:
    publication_info = dict(spec.get("publication_info") or {})
    if spec.get("paper_link") and not publication_info.get("paper_link"):
        publication_info["paper_link"] = spec.get("paper_link")

    if spec.get("include_required_info", True):
        for key, label in PUBLICATION_FIELDS:
            value = str(publication_info.get(key) or "").strip()
            add_text_paragraph(doc, f"{label}：{value}")
        return

    paper_link = str(publication_info.get("paper_link") or "").strip()
    if paper_link:
        add_text_paragraph(doc, "论文链接：" + paper_link)


def build(spec: dict[str, Any], output: Path, template: Path | None = None) -> None:
    if template and template.exists():
        doc = Document(str(template))
        clear_document(doc)
    else:
        doc = Document()
    configure_document(doc)

    title = (spec.get("title") or "").strip()
    if not title:
        raise ValueError("Spec must include a non-empty title.")
    add_text_paragraph(doc, title, align=WD_ALIGN_PARAGRAPH.CENTER, size_pt=float(spec.get("title_size_pt") or 14), bold=bool(spec.get("title_bold", False)))

    for paragraph in spec.get("paragraphs") or []:
        text = str(paragraph).strip()
        if text:
            add_text_paragraph(doc, text)

    add_publication_info(doc, spec)

    figures = spec.get("figures") or []
    if figures and spec.get("page_break_before_figures", False):
        doc.add_page_break()

    for figure in figures:
        add_figure(doc, figure, output.parent)

    output.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output))


def main() -> int:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("spec", type=Path, help="JSON file with title, paragraphs, publication_info fixed fields, and figures.")
    parser.add_argument("output", type=Path)
    parser.add_argument("--template", type=Path)
    args = parser.parse_args()

    spec = json.loads(args.spec.read_text(encoding="utf-8"))
    build(spec, args.output, args.template)
    print(str(args.output))
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
